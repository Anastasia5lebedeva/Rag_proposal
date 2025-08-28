from __future__ import annotations

import argparse
import sys
from pathlib import Path
import re

from docx import Document

from .logger import get_logger

log = get_logger(__name__)

HEADING_RX = re.compile(r"^(heading|zagolovok)\s*(\d+)?", re.IGNORECASE)
DEFAULT_MAX_PARAS = 5
DEFAULT_MAX_ROWS = 3


def heading_level(par):
    name = getattr(getattr(par, "style", None), "name", "") or ""
    m = HEADING_RX.match(name)
    if m:
        try:
            return int(m.group(2) or 1)
        except Exception:  # noqa: BLE001
            return 1
    return 0


def preview_docx(path: Path, max_paras: int = 5, max_rows: int = 3) -> int:
    if not path.exists():
        log.error("file not found: %s", path)
        return 1
    try:
        doc = Document(str(path))
    except Exception as e:  # noqa: BLE001
        log.error("cannot open DOCX: %s", e)
        return 2

    paras = list(doc.paragraphs)
    tables = list(doc.tables)
    headings = sum(1 for x in paras if heading_level(x) > 0)
    log.info("paras: %s tables: %s headings: %s", len(paras), len(tables), headings)

    for x in paras[:max_paras]:
        name = getattr(getattr(x, "style", None), "name", "") or ""
        log.info("paragraph: %s | style: %s", x.text.strip(), name)

    if tables:
        for ti, t in enumerate(tables, 1):
            log.info("table %s", ti)
            for r in t.rows[:max_rows]:
                log.info("row: %s", "|".join(c.text.strip() for c in r.cells))
    return 0


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect docx structure")
    parser.add_argument("path", help="path to DOCX file")
    parser.add_argument("--max-paras", type=int, default=DEFAULT_MAX_PARAS)
    parser.add_argument("--max-rows", type=int, default=DEFAULT_MAX_ROWS)
    args = parser.parse_args()
    code = preview_docx(Path(args.path).expanduser(), args.max_paras, args.max_rows)
    sys.exit(code)


if __name__ == "__main__":
    main()